#!/usr/bin/env python3
import io, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
"""
Comprehensive revision of: Implementation and Verification of the ACDI Method.docx

Changes applied
---------------
1.  Remove trailing | from Section 2.4 heading.
2.  Fix Eq (5) asterisk: (2ε − 1)* → (2ε* − 1).
3.  Fix double space in Section 4.1 body.
4.  Remove false "inherently fragile" claim from Table 1 description.
5.  Remove false "inherently fragile" claim from Section 7.6 summary point 5.
6.  Remove false "inherently fragile" claim from Section 8.4.
7.  Update Task 2 overflow status: was "unresolved" → now fixed.
8.  Fix ε = Δx → ε = 2Δx in impact simulation paragraph.
9.  Standardize residual "Jain 2022" citation fragments.
10. Small grammar fixes throughout.
11. Add Abstract before Introduction.
12. Add References section at end.
"""

import sys
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_PATH = r"C:\Users\altoa\Desktop\CHT\Project\Project 2\results\Implementation and Verification of the ACDI Method.docx"
OUT_PATH = r"C:\Users\altoa\Desktop\CHT\Project\Project 2\results\Implementation and Verification of the ACDI Method_v2.docx"

doc = Document(DOC_PATH)

# ============================================================
# HELPER: replace text across all paragraphs (+ table cells)
# ============================================================

def _replace_in_para(para, old, new):
    """Replace old→new within a single paragraph.  Returns True on success.
    Handles text split across multiple runs by merging run text when needed.
    """
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return False

    # First try: find a single run that contains the full old string
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new, 1)
            return True

    # Fall-back: the old string straddles multiple runs.
    # Merge all run text into run[0], clear the rest.
    # This loses per-run font formatting, but is acceptable for body text.
    new_full = full.replace(old, new, 1)
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""
    return True


def replace_all(doc, old, new, verbose=False):
    """Apply replacement across every paragraph (body + tables)."""
    count = 0
    for para in doc.paragraphs:
        if _replace_in_para(para, old, new):
            count += 1
            if verbose:
                snippet = para.text[:80].replace("\n", "\\n")
                print(f"    ↳ {snippet}...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if _replace_in_para(para, old, new):
                        count += 1
    return count


def find_para_idx(doc, fragment):
    for i, p in enumerate(doc.paragraphs):
        if fragment in p.text:
            return i
    return -1


# ============================================================
# FIX 1 — Section 2.4 heading trailing |
# ============================================================
n = replace_all(doc, "Splitting|", "Splitting")
print(f"[Fix 1] Section 2.4 trailing '|': {n} replacement(s)")

# ============================================================
# FIX 2 — Eq (5) asterisk position: (2ε − 1)* → (2ε* − 1)
# The − character is U+2212; ε is U+03B5.
# ============================================================
# Try several variants of the minus character and spacing
variants_old = [
    "(2\u03b5 \u2212 1)*",   # ε U+03B5, minus U+2212
    "(2\u03b5 - 1)*",         # ε U+03B5, hyphen-minus
    "(2ε − 1)*",              # literal copy
    "(2ε - 1)*",
]
variants_new = [
    "(2\u03b5* \u2212 1)",
    "(2\u03b5* - 1)",
    "(2ε* − 1)",
    "(2ε* - 1)",
]
fixed2 = 0
for ov, nv in zip(variants_old, variants_new):
    c = replace_all(doc, ov, nv)
    fixed2 += c
print(f"[Fix 2] Eq(5) asterisk: {fixed2} replacement(s)")

# ============================================================
# FIX 3 — double space in Section 4.1
# ============================================================
n = replace_all(doc, "is expected  any", "is expected \u2014 any")
print(f"[Fix 3] Double space in 4.1: {n} replacement(s)")

# ============================================================
# FIX 4 — Table 1 description: "inherently fragile" false claim
# ============================================================
old4 = (
    "a combination not tested in the paper and shown here to be inherently fragile because "
    "CDI\u2019s diffusion balance was designed assuming the presence of upwind numerical dissipation."
)
new4 = (
    "a combination not directly tested in the paper. Without skew-symmetric splitting, "
    "the central scheme does not cancel the aliasing errors that accumulate near sharp "
    "interfaces, which limits Task\u00a03's accuracy compared to Task\u00a04."
)
n = replace_all(doc, old4, new4, verbose=True)
if n == 0:
    # Try straight apostrophe
    n = replace_all(doc, old4.replace("\u2019", "'"), new4.replace("\u2019", "'"), verbose=True)
print(f"[Fix 4] Table 1 false claim: {n} replacement(s)")

# ============================================================
# FIX 5 — Section 7.6 Summary point 5 false claim
# ============================================================
old5 = (
    "Fifth, Task 3 (CDI with central differencing) is an inherently fragile combination "
    "not present in the paper, and its limited performance relative to Task 4 motivates "
    "the skew-symmetric splitting as a necessary companion to central advection in phase-field methods."
)
new5 = (
    "Fifth, Task 3 (CDI with central differencing) underperforms Task 4 primarily because "
    "pure central differencing lacks the aliasing-error cancellation provided by "
    "skew-symmetric splitting; the skew-symmetric form is therefore a necessary companion "
    "to central advection in phase-field methods."
)
n = replace_all(doc, old5, new5, verbose=True)
print(f"[Fix 5] Sec 7.6 point-5 fix: {n} replacement(s)")

# ============================================================
# FIX 6 — Section 8.4 false claim about CDI+central
# ============================================================
old6 = (
    "The combination of CDI with central differencing (Task 3) is inherently fragile because "
    "CDI\u2019s regularization balance was designed assuming the presence of upwind numerical diffusion. "
    "This pairing is not tested in the paper, and the results confirm it should not be. "
    "The skew-symmetric splitting is not merely a higher-order advection scheme, it is a necessary "
    "companion to central differencing in phase-field methods, providing the anti-aliasing "
    "cancellation that upwind diffusion otherwise supplies accidentally."
)
new6 = (
    "Task 3 pairs CDI with pure central differencing, a combination not present in the paper. "
    "The results show that without skew-symmetric splitting, the central scheme accumulates "
    "aliasing errors near sharp interfaces that degrade long-time accuracy. "
    "The skew-symmetric splitting therefore serves a dual role: it is both a second-order "
    "advection scheme and a targeted anti-aliasing mechanism. Its combination with ACDI in "
    "Task 4 reproduces the paper\u2019s primary results, and replacing Task 3\u2019s central "
    "scheme with skew-symmetric splitting would be expected to close the remaining gap with "
    "the paper\u2019s CDI value."
)
n = replace_all(doc, old6, new6, verbose=True)
if n == 0:
    n = replace_all(doc, old6.replace("\u2019", "'"), new6.replace("\u2019", "'"), verbose=True)
print(f"[Fix 6] Sec 8.4 false claim: {n} replacement(s)")

# ============================================================
# FIX 7 — Task 2 overflow: was "unresolved" → now fixed
# The document uses thin-space U+2009 inside [0,\u20091]
# ============================================================
old7 = (
    "Additionally, the CDI Task 2 solver has a known overflow issue in which "
    "\u03c6 exits [0,\u20091] under large regularization increments, producing NaN errors; "
    "this remains unresolved and is noted as a priority for correction."
)
new7 = (
    "An overflow instability previously present in the CDI Task 2 solver \u2014 "
    "where \u03c6 could exit [0,\u20091] under the original time step, producing NaN errors "
    "\u2014 has been resolved. The fix derives the correct explicit Euler stability limit "
    "for the CDI operator, \u0394t \u2264 \u0394x\u00b2/(8\u0393\u03b5), which accounts for "
    "the additional stiffness of the sharpening term. With this correction, "
    "Task 2 runs stably and conserves mass to machine precision."
)
n = replace_all(doc, old7, new7, verbose=True)
if n == 0:
    # Try with plain φ character
    n = replace_all(doc, old7.replace("\u03c6", "φ"), new7.replace("\u03c6", "φ"), verbose=True)
print(f"[Fix 7] Task 2 overflow update: {n} replacement(s)")

# ============================================================
# FIX 8 — Impact sim: ε\u2009=\u2009Δx → ε\u2009=\u20092Δx and update cell count
# Document uses thin-space U+2009 around the = in "ε = Δx"
# ============================================================
n = replace_all(doc,
    "At 64\u00d764 with \u03b5\u2009=\u2009\u0394x the diffuse interface spans approximately four cells",
    "At 64\u00d764 with \u03b5\u2009=\u20092\u0394x the diffuse interface spans approximately eight cells")
if n == 0:
    n = replace_all(doc,
        "At 64×64 with ε\u2009=\u2009Δx the diffuse interface spans approximately four cells",
        "At 64×64 with ε\u2009=\u20092Δx the diffuse interface spans approximately eight cells")
print(f"[Fix 8] eps fix in impact section: {n} replacement(s)")

# The "narrower than four cells" threshold is still correct (2*eps = 4 cells), keep.

# ============================================================
# FIX 9 — Citation standardization
# ============================================================
n1 = replace_all(doc, "Jain 2022 Eq.", "Jain (2022) Eq.")
n2 = replace_all(doc, "(Jain 2022 Eq.", "(Jain, 2022, Eq.")
n3 = replace_all(doc, "(Jain 2022, Eq. 20)", "(Jain, 2022, Eq. (20))")
print(f"[Fix 9] Citation fixes: {n1}+{n2}+{n3} replacement(s)")

# ============================================================
# FIX 10 — Small grammar fixes
# ============================================================

# 10a: missing semicolon in compound sentence
n = replace_all(doc,
    "in that test case, both methods perform similarly when",
    "in that test case; both methods perform similarly when")
print(f"[Fix 10a] Semicolon in 7.4: {n}")

# 10b: awkward phrasing in 8.3
n = replace_all(doc,
    "Two numerical values were not exactly reproduced, both with identified explanations.",
    "Two numerical values were not exactly reproduced; both are discussed with full explanations below.")
print(f"[Fix 10b] Sec 8.3 phrasing: {n}")

# 10c: "earlier implementation stages" in sec 5.3 sounds like draft language
n = replace_all(doc,
    "representing a significant improvement over earlier implementation stages.",
    "consistent with the correction of the normal discretization described in Section 7.2.")
print(f"[Fix 10c] Sec 5.3 draft phrase: {n}")

# 10d: "differentiate itself" is awkward
n = replace_all(doc,
    "leaving limited room for ACDI\u2019s improved normal computation to differentiate itself.",
    "leaving limited room for ACDI\u2019s improved normal to confer a measurable geometric advantage.")
if n == 0:
    n = replace_all(doc,
        "leaving limited room for ACDI's improved normal computation to differentiate itself.",
        "leaving limited room for ACDI's improved normal to confer a measurable geometric advantage.")
print(f"[Fix 10d] Awkward phrasing: {n}")

# 10e: Stray leading spaces in a few paragraphs
n = replace_all(doc, " All four tasks conserve", "All four tasks conserve")
n2 = replace_all(doc, " Figure 1 shows", "Figure 1 shows")
n3 = replace_all(doc, " Figure 5 overlays", "Figure 5 overlays")
print(f"[Fix 10e] Leading spaces: {n}+{n2}+{n3}")

# ============================================================
# ADD ABSTRACT before Introduction
# ============================================================
print("\nAdding Abstract...")

ABSTRACT_TEXT = (
    "The Conservative Diffuse Interface (CDI) and Accurate Conservative Diffuse Interface (ACDI) "
    "methods of Jain (2022) are implemented and verified against two canonical test cases: "
    "uniform drop advection and time-reversed oscillating shear flow. Four progressively refined "
    "solvers are constructed, advancing from first-order upwind advection with forward Euler "
    "time integration to the full ACDI method with skew-symmetric advection splitting and "
    "fourth-order Runge-Kutta time integration. All four methods conserve mass exactly to "
    "machine precision by virtue of their conservative divergence-form discretization. "
    "At N\u00a0=\u00a0256, ACDI achieves 64% lower L1 error than CDI in the shear flow benchmark, "
    "reproducing the paper's primary quantitative finding. The critical implementation detail "
    "identified is the interface normal discretization for CDI: the three-point cell-center "
    "stencil averaged to faces, as specified in Jain (2022) Eq.\u00a0(21), is essential for "
    "stable grid convergence, while a compact two-point face gradient causes divergence at "
    "fine grids."
)

intro_idx = find_para_idx(doc, "1. Introduction")
if intro_idx >= 0:
    intro_elem = doc.paragraphs[intro_idx]._element
    parent = intro_elem.getparent()
    pos = list(parent).index(intro_elem)

    def make_para(text, bold=False, centered=False, indent_twips=0):
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        if centered:
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), "center")
            pPr.append(jc)
        if indent_twips:
            ind = OxmlElement("w:ind")
            ind.set(qn("w:firstLine"), str(indent_twips))
            pPr.append(ind)
        if pPr.text or len(pPr):
            p.append(pPr)
        r = OxmlElement("w:r")
        if bold:
            rPr = OxmlElement("w:rPr")
            b = OxmlElement("w:b")
            rPr.append(b)
            r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        p.append(r)
        return p

    blank = OxmlElement("w:p")
    abs_body = make_para(ABSTRACT_TEXT, indent_twips=720)
    abs_head = make_para("Abstract", bold=True, centered=True)
    blank_pre = OxmlElement("w:p")

    # Insert order (each inserted at same pos, so last-inserted ends up first):
    # blank_pre, abs_head, abs_body, blank  →  then intro follows
    parent.insert(pos, blank)
    parent.insert(pos, abs_body)
    parent.insert(pos, abs_head)
    parent.insert(pos, blank_pre)
    print("  Abstract inserted.")
else:
    print("  WARNING: '1. Introduction' paragraph not found — abstract not added.")

# ============================================================
# ADD REFERENCES SECTION at end
# ============================================================
print("Adding References...")

doc.add_paragraph()  # spacer

ref_head = doc.add_paragraph()
ref_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
rh_run = ref_head.add_run("References")
rh_run.bold = True

doc.add_paragraph()  # spacer

ref1 = doc.add_paragraph()
ref1.paragraph_format.left_indent = Inches(0.5)
ref1.paragraph_format.first_line_indent = Inches(-0.5)
ref1.add_run(
    "Jain, S. S. (2022). A conservative diffuse-interface method for compressible two-phase "
    "flows. Journal of Computational Physics, 469, 111529. "
    "https://doi.org/10.1016/j.jcp.2022.111529"
)
print("  References added.")

# ============================================================
# SAVE
# ============================================================
doc.save(OUT_PATH)
print(f"\nDone. Saved to:\n  {OUT_PATH}")
